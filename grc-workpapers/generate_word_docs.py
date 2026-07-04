"""
Generates Big 4 style Word deliverables for the GRC / Oracle ITGC-ITAC engagement:
  1. Scoping_Memorandum.docx
  2. Findings_Report.docx
For education / template use - replace sample data with actual engagement data.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x59, 0x59, 0x59)
RED = RGBColor(0xC0, 0x00, 0x00)
AMBER = RGBColor(0xBF, 0x8F, 0x00)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def base_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)


def add_cover(doc, title, subtitle):
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(subtitle)
    sr.font.size = Pt(12)
    sr.font.color.rgb = GREY
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("Client: ABC Manufacturing Ltd.   |   Period: FY 2025-26   |   System: Oracle Fusion Cloud ERP (24C)")
    mr.font.size = Pt(10)
    mr.font.color.rgb = GREY
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("Template for educational use - replace sample data with actual engagement data.")
    nr.italic = True
    nr.font.size = Pt(8.5)
    nr.font.color.rgb = GREY


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = NAVY
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    return p


def kv_table(doc, rows):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in rows:
        cells = tbl.add_row().cells
        cells[0].width = Inches(2.2)
        rk = cells[0].paragraphs[0].add_run(k)
        rk.bold = True
        cells[1].paragraphs[0].add_run(v)
    return tbl


def grid_table(doc, headers, rows, col_widths=None, sev_col=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], "2E5496")
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            if sev_col is not None and i == sev_col:
                if "Material Weakness" in str(val):
                    set_cell_bg(cells[i], "FFC7CE"); run.bold = True; run.font.color.rgb = RED
                elif "Significant" in str(val):
                    set_cell_bg(cells[i], "FFEB9C"); run.font.color.rgb = AMBER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    return tbl


# =====================================================================
# DOC 1: SCOPING MEMORANDUM
# =====================================================================
doc = Document()
base_styles(doc)
add_cover(doc, "Scoping Memorandum", "ITGC & ITAC Audit - Oracle Fusion Cloud ERP")
doc.add_paragraph()

h1(doc, "1. Engagement Details")
kv_table(doc, [
    ("Client", "ABC Manufacturing Ltd."),
    ("Period under review", "01-Apr-2025 to 31-Mar-2026 (FY 2025-26)"),
    ("System", "Oracle Fusion Cloud ERP (Release 24C)"),
    ("Prepared by", "________________    Date: __/__/____"),
    ("Reviewed by", "________________    Date: __/__/____"),
])

h1(doc, "2. Objective")
doc.add_paragraph(
    "To evaluate the design and operating effectiveness of IT General Controls (ITGC) and IT "
    "Application Controls (ITAC) over financially significant Oracle Fusion applications supporting "
    "Internal Control over Financial Reporting (ICFR) / SOX Section 404 compliance.")

h1(doc, "3. Scoping Approach")
doc.add_paragraph(
    "Applications and processes are scoped based on financial materiality, relevance to significant "
    "accounts and disclosures, and risk of material misstatement. A controls-reliance strategy is "
    "adopted, supplemented by limited substantive testing where required.")

h1(doc, "4. In-scope Applications")
grid_table(doc, ["Application / Module", "Significant?", "Basis"], [
    ["Oracle General Ledger", "Yes", "Financial statement close"],
    ["Oracle Payables", "Yes", "Expenses / Accounts Payable"],
    ["Oracle Receivables", "Yes", "Revenue / Accounts Receivable"],
    ["Oracle Fixed Assets", "Yes", "Depreciation / PP&E"],
    ["Oracle Inventory", "Yes", "Cost of Goods Sold"],
    ["Oracle HCM Payroll", "Yes", "Payroll expense"],
], col_widths=[3.0, 1.2, 2.5])

h1(doc, "5. In-scope Business Processes")
for line in [
    "Procure-to-Pay (P2P) - requisition to payment",
    "Order-to-Cash (O2C) - order to collection",
    "Record-to-Report (R2R) - journal to financial reporting",
    "Inventory - receipt to valuation",
    "Payroll (Hire-to-Retire) - onboarding to disbursement",
]:
    doc.add_paragraph(line, style="List Bullet")

h1(doc, "6. Control Domains Tested")
h2(doc, "6.1 ITGC")
for line in ["Access to Programs & Data", "Change Management", "IT Operations (batch/interfaces)", "Backup & Recovery"]:
    doc.add_paragraph(line, style="List Bullet")
h2(doc, "6.2 ITAC")
for line in ["Automated matching (e.g., 3-way match)", "Automated calculations", "System validations & edit checks", "Automated approval workflows"]:
    doc.add_paragraph(line, style="List Bullet")

h1(doc, "7. Materiality")
kv_table(doc, [
    ("Overall materiality", "INR 12.5 Cr"),
    ("Performance materiality", "INR 8.0 Cr"),
    ("Clearly trivial threshold", "INR 0.6 Cr"),
])

h1(doc, "8. Testing Approach & Timeline")
doc.add_paragraph(
    "Design effectiveness will be assessed via walkthroughs and configuration review. Operating "
    "effectiveness will be assessed through attribute-based sampling for manual/IT-dependent controls "
    "and test-of-one plus ITGC reliance for automated controls. Interim testing followed by roll-forward "
    "to period end.")

h1(doc, "9. Reliance & Dependencies")
doc.add_paragraph(
    "Reliance on automated (ITAC) controls is dependent on effective Change Management ITGC. Any report "
    "used in testing (IPE) will have its completeness and accuracy validated before use.")

doc.save("/projects/sandbox/Vivek-Kriplani/grc-workpapers/Scoping_Memorandum.docx")
print("Saved Scoping_Memorandum.docx")


# =====================================================================
# DOC 2: FINDINGS REPORT
# =====================================================================
doc = Document()
base_styles(doc)
add_cover(doc, "Findings Report", "ITGC & ITAC Audit - Summary of Deficiencies")
doc.add_paragraph()

h1(doc, "1. Executive Summary")
doc.add_paragraph(
    "This report summarises control deficiencies identified during the ITGC and ITAC audit of Oracle "
    "Fusion Cloud ERP for FY 2025-26. Findings are rated by severity and accompanied by root cause, "
    "impact, recommendation and management response. One Material Weakness, one Significant Deficiency "
    "and two Control Deficiencies were identified.")

h1(doc, "2. Rating Summary")
grid_table(doc, ["Rating", "Count", "Definition (summary)"], [
    ["Material Weakness", "1", "Reasonable possibility that a material misstatement is not prevented/detected"],
    ["Significant Deficiency", "1", "Important enough to merit governance attention; not material"],
    ["Control Deficiency", "2", "Control missing or not operating as designed"],
], col_widths=[2.0, 0.8, 4.0], sev_col=0)

h1(doc, "3. Detailed Findings")

findings = [
    {
        "num": "D-08",
        "title": "Excessive privileged access - SoD violation",
        "rating": "Material Weakness",
        "condition": "Three finance users hold both User Administration and transaction posting access in Oracle Fusion.",
        "criteria": "Segregation of Duties principle - access administration must be segregated from transaction processing (Company Policy SEC-04 / SOX).",
        "cause": "Roles were granted during system migration and never subsequently reviewed.",
        "effect": "Risk of unauthorized access grant and fraudulent entries; potential material misstatement.",
        "rec": "Remove the User Administration duty from finance roles and implement quarterly access certification.",
        "resp": "Agreed. Owner: IT Security Head. Target date: 15-Jul-2026.",
    },
    {
        "num": "D-05",
        "title": "Inadequate change management segregation",
        "rating": "Significant Deficiency",
        "condition": "A developer had access to migrate changes directly to the production environment.",
        "criteria": "Developers must be segregated from production migration per the change management policy.",
        "cause": "No enforced segregation within the deployment process.",
        "effect": "Unauthorized or untested changes could reach production undetected.",
        "rec": "Segregate developer and migration roles; enforce a migration approval gate.",
        "resp": "Agreed. Owner: IT Manager. Target date: 31-Aug-2026.",
    },
    {
        "num": "D-03",
        "title": "User provisioned without documented approval",
        "rating": "Control Deficiency",
        "condition": "1 of 25 sampled new users was provisioned without evidence of approval.",
        "criteria": "New user access must be authorized before provisioning.",
        "cause": "Manual provisioning process without an enforced approval workflow.",
        "effect": "Risk of inappropriate access being granted.",
        "rec": "Implement an enforced access-request workflow prior to grant.",
        "resp": "Agreed. Owner: HR + IT. Target date: 30-Sep-2026.",
    },
    {
        "num": "D-11",
        "title": "Depreciation rounding configuration difference",
        "rating": "Control Deficiency",
        "condition": "One sampled asset showed an immaterial depreciation rounding difference (INR 300).",
        "criteria": "Depreciation must be calculated accurately per policy.",
        "cause": "Rounding configuration setting.",
        "effect": "Immaterial; no impact on financial statements.",
        "rec": "Adjust rounding configuration; monitor via GL reconciliation.",
        "resp": "Agreed. Owner: Finance Systems. Target date: 30-Sep-2026.",
    },
]

for f in findings:
    hp = doc.add_paragraph()
    hr = hp.add_run("Finding %s: %s" % (f["num"], f["title"]))
    hr.bold = True
    hr.font.size = Pt(11.5)
    hr.font.color.rgb = NAVY
    grid_table(doc, ["Attribute", "Description"], [
        ["Rating", f["rating"]],
        ["Condition", f["condition"]],
        ["Criteria", f["criteria"]],
        ["Cause", f["cause"]],
        ["Effect", f["effect"]],
        ["Recommendation", f["rec"]],
        ["Management Response", f["resp"]],
    ], col_widths=[1.8, 5.0], sev_col=1)
    doc.add_paragraph()

h1(doc, "4. Remediation Tracking")
doc.add_paragraph(
    "Each finding is assigned an owner and target date and tracked to closure. Remediation actions will "
    "be re-tested by the audit team; evidence obtained will be re-performed before a finding is closed. "
    "Refer to the Remediation Roadmap tab in the workpaper package for the phased action plan.")

doc.save("/projects/sandbox/Vivek-Kriplani/grc-workpapers/Findings_Report.docx")
print("Saved Findings_Report.docx")
