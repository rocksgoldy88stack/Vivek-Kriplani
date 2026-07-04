"""
Generates a detailed Word deliverable:
  AAC_Access_Model_Config_Guide.docx
Step-by-step Oracle Fusion Advanced Access Controls (AAC) Access Model build & deploy guide.
For education / template use. Oracle screen labels/navigation may vary slightly by release.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x59, 0x59, 0x59)
GREEN = RGBColor(0x37, 0x56, 0x23)
RED = RGBColor(0xC0, 0x00, 0x00)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def base_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(10.5)


def add_cover(doc, title, subtitle):
    doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title); run.bold = True; run.font.size = Pt(22); run.font.color.rgb = NAVY
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(subtitle); sr.font.size = Pt(12); sr.font.color.rgb = GREY
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("Oracle Risk Management Cloud - Advanced Access Controls (AAC)")
    mr.font.size = Pt(10); mr.font.color.rgb = GREY
    note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("Template for educational use. Screen labels & navigation may vary by Oracle release.")
    nr.italic = True; nr.font.size = Pt(8.5); nr.font.color.rgb = GREY


def h1(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY
    return p


def h2(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
    return p


def para(doc, text):
    return doc.add_paragraph(text)


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def step(doc, num, text):
    p = doc.add_paragraph()
    r = p.add_run("Step %s: " % num); r.bold = True; r.font.color.rgb = NAVY
    p.add_run(text)
    return p


def screen(doc, label):
    """A 'Screen / Navigation' callout line."""
    p = doc.add_paragraph()
    r = p.add_run("[Screen] "); r.bold = True; r.font.color.rgb = GREEN; r.font.size = Pt(9.5)
    rr = p.add_run(label); rr.font.size = Pt(9.5); rr.font.color.rgb = GREY; rr.italic = True
    return p


def grid_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers)); tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], "2E5496")
        run = hdr[i].paragraphs[0].add_run(h); run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); run.font.size = Pt(9.5)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(val)); run.font.size = Pt(9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    return tbl


doc = Document()
base_styles(doc)
add_cover(doc, "AAC Access Model Configuration Guide",
          "Step-by-Step: Build, Deploy & Monitor SoD Access Controls")
doc.add_paragraph()

# ---------------------------------------------------------
h1(doc, "1. Purpose & Scope")
para(doc,
     "This guide provides step-by-step instructions to build a Segregation of Duties (SoD) "
     "Access Model in Oracle Advanced Access Controls (AAC), deploy it as an Access Control for "
     "continuous monitoring, and manage the resulting incidents. It is intended for GRC analysts "
     "and Oracle security administrators.")

# ---------------------------------------------------------
h1(doc, "2. Key Concepts")
grid_table(doc, ["Term", "Meaning"], [
    ["Access Point", "A grantable security artefact - a privilege, aggregate privilege, duty role, job role or function that lets a user perform an action."],
    ["Entitlement", "A reusable, named group of related access points, used to keep models consistent."],
    ["Access Model", "The design/simulation object where conflict logic is defined and analysed (like a sandbox). No incidents are generated for records here."],
    ["Access Control", "A deployed model that runs against live data on each synchronisation and generates incidents."],
    ["Incident", "A record showing a user who violates the control (holds the conflicting access)."],
    ["Path-based analysis", "AAC traces the full security path (User -> Role -> Duty -> Privilege) to confirm access is actually reachable, reducing false positives."],
])
para(doc, "Rule of thumb: Model = blueprint you test and tune; Control = the running engine that monitors and raises incidents.")

# ---------------------------------------------------------
h1(doc, "3. Prerequisites")
bullet(doc, "Roles: an AAC-enabled role such as those granting 'Manage Models' and 'Manage Controls' (e.g., Access Model/Control management duties within the Risk Management job roles).")
bullet(doc, "Global User Synchronization / data sync jobs must have completed so AAC has current user, role and access-point data.")
bullet(doc, "Perspectives (optional) set up for organising and routing controls (e.g., by Business Unit or Process).")
bullet(doc, "An agreed SoD ruleset (see 12_SoD_Ruleset tab in the workpaper package) defining the conflicting function pairs.")

# ---------------------------------------------------------
h1(doc, "4. Navigation")
screen(doc, "Navigator (hamburger menu) > Risk Management > Advanced Controls")
para(doc, "The Advanced Controls work area opens with tabs including Models, Controls, Results, and Scheduling. All model-building happens under the Models tab.")

# ---------------------------------------------------------
h1(doc, "5. Build an Access Model (Step-by-Step)")
para(doc, "Worked example: build a model for SOD-01 - 'Maintain Supplier' conflicts with 'Process AP Payment'.")

step(doc, 1, "Open the Models tab.")
screen(doc, "Advanced Controls > Models")

step(doc, 2, "Create a new access model.")
screen(doc, "Actions menu > Create > Access Model")
bullet(doc, "Name: SOD-01 Maintain Supplier vs Process Payment")
bullet(doc, "Description: Detects users who can both maintain suppliers and process AP payments.")
bullet(doc, "State/Status: set to Active (or Draft while still building).")

step(doc, 3, "Add the first access point group (Function A).")
screen(doc, "Model definition page > Access Point Groups / Add Filter")
bullet(doc, "Create a group and add the access point(s) representing 'Maintain Supplier' (e.g., the privilege/duty 'Manage Suppliers' / 'Supplier Profile Management Duty').")
bullet(doc, "Prefer adding an Entitlement if one exists, so the same access-point set is reused across models.")

step(doc, 4, "Add the second access point group (Function B).")
bullet(doc, "Create a second group and add the access point(s) representing 'Process AP Payment' (e.g., 'Create Payment' / 'Payables Payment Duty').")

step(doc, 5, "Define the conflict logic between the groups.")
bullet(doc, "Set the access-within logic (user has ANY vs ALL of the access points in a group) - typically ANY for each group.")
bullet(doc, "Set the relationship so an incident is raised when a user has access in Group A AND access in Group B.")

step(doc, 6, "Add condition filters to reduce noise (optional but recommended).")
screen(doc, "Model definition page > Conditions / Filters")
bullet(doc, "Exclude legitimate exceptions - e.g., exclude specific IT admin users, or filter by an attribute such as active users only.")
bullet(doc, "You can filter by user, role, or access-point attributes.")

step(doc, 7, "Analyse the model.")
screen(doc, "Actions menu > Analyze (results preview)")
bullet(doc, "AAC runs the logic and shows the count of users/records that would violate. Review the sample results.")
bullet(doc, "If too many false positives appear, refine conditions (Step 6) and re-analyse. Iterate until results are meaningful.")

step(doc, 8, "Save the model.")
bullet(doc, "Save once the analysed results look accurate. The model is now ready to deploy as a control.")

para(doc, "Repeat Steps 1-8 for each rule in the SoD ruleset (or import multiple models where your process allows).")

# ---------------------------------------------------------
h1(doc, "6. Deploy the Model as an Access Control")
step(doc, 1, "Select the saved model and deploy it.")
screen(doc, "Models > select model > Actions > Deploy as Control (or Deploy)")

step(doc, 2, "Complete the control attributes.")
grid_table(doc, ["Attribute", "Guidance"], [
    ["Name / Description", "Carry over a clear, business-readable name."],
    ["Priority", "Set High/Medium/Low aligned to the rule severity in the SoD ruleset."],
    ["Result Type", "Access incidents (users who violate)."],
    ["Status", "Active (so it runs on each sync)."],
    ["Perspectives", "Assign perspective values (e.g., Process = P2P, BU = India) for routing & reporting."],
])

step(doc, 3, "Assign result investigators / owners.")
bullet(doc, "Assign the users or groups who will investigate and disposition incidents (e.g., GRC analyst, control owner).")

step(doc, 4, "Confirm the run schedule.")
screen(doc, "Advanced Controls > Scheduling")
bullet(doc, "Controls evaluate when the synchronization/analysis job runs. Ensure a regular schedule (e.g., daily/weekly) is active.")

step(doc, 5, "Submit / activate the control.")
bullet(doc, "On the next scheduled run, the control generates incidents for all users who hold the conflicting access across the full population.")

# ---------------------------------------------------------
h1(doc, "7. Manage Incidents (Results)")
screen(doc, "Advanced Controls > Results > Access")
para(doc, "Each incident represents a user violating the control. Investigate and set a disposition using the incident status workflow.")
grid_table(doc, ["Incident Status", "Meaning / When to use"], [
    ["Assigned", "Newly generated, awaiting investigation."],
    ["In Investigation", "Analyst is reviewing the conflict and business context."],
    ["Accepted", "Risk formally accepted with documented justification (management sign-off)."],
    ["Remediate", "Access will be removed / role redesigned to clear the conflict."],
    ["Resolved / Closed", "Conflict addressed; on next run it no longer appears (Closed)."],
    ["Control Inactive", "Set when a control is retired; associated incidents are closed."],
])
para(doc, "For each incident, record a comment, attach the mitigating/compensating control where access must remain, and route to the owner. Track remediation in the 17_Remediation_Roadmap tab.")

# ---------------------------------------------------------
h1(doc, "8. Tuning & Best Practices")
bullet(doc, "Start broad, then tune: expect many results on first run; refine conditions to cut false positives (target < 10%).")
bullet(doc, "Use Entitlements for reusable access-point groups so models stay consistent and maintainable.")
bullet(doc, "Rely on path-based analysis - confirm access is actually reachable, not just assigned.")
bullet(doc, "Use global conditions to consistently exclude known legitimate access (e.g., break-glass IT admin).")
bullet(doc, "Never deploy seeded roles unchanged - copy, then remove conflicting duties (see 14_Role_Design tab).")
bullet(doc, "Simulate role changes in a model before provisioning to a user (preventive SoD).")
bullet(doc, "Schedule regular synchronization so monitoring stays near-real-time.")
bullet(doc, "Review dashboard KPIs each cycle (see 19_Monitoring_Dashboard tab) - aging, MTTR, false-positive rate.")

# ---------------------------------------------------------
h1(doc, "9. Common Conflict Models (starter set)")
grid_table(doc, ["Rule", "Group A (Function)", "Group B (Conflicts with)", "Priority"], [
    ["SOD-01", "Maintain Supplier", "Process AP Payment", "High"],
    ["SOD-03", "Enter Purchase Order", "Approve Purchase Order", "High"],
    ["SOD-28", "Create Journal Entry", "Approve Journal Entry", "High"],
    ["SOD-65", "Maintain Bank Account", "Make Payment", "High"],
    ["SOD-73", "User Administration", "Any Financial Transaction", "Critical"],
], col_widths=[0.8, 2.3, 2.3, 1.0])
para(doc, "Full 87-rule ruleset is in the 12_SoD_Ruleset tab of the workpaper package.")

# ---------------------------------------------------------
h1(doc, "10. Troubleshooting")
grid_table(doc, ["Symptom", "Likely Cause", "Action"], [
    ["Model returns 0 results but conflicts exist", "Wrong/too-narrow access points; user sync not run", "Verify access points map to real privileges; run synchronization; re-analyse."],
    ["Too many false positives", "Logic too broad; no conditions", "Add condition filters; use path-based analysis; exclude known-good access."],
    ["Incident not clearing after fix", "Sync/control not re-run", "Re-run synchronization; confirm access actually removed; status moves to Closed."],
    ["Legit admin flagged repeatedly", "No exclusion condition", "Add a global condition/exclusion with documented justification (mitigating control)."],
])

doc.save("/projects/sandbox/Vivek-Kriplani/grc-workpapers/AAC_Access_Model_Config_Guide.docx")
print("Saved AAC_Access_Model_Config_Guide.docx")
